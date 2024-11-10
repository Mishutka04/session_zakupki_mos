import re
import os
import asyncio
import httpx  # для асинхронных запросов
from docx import Document
import fitz  # для обработки PDF
from gradio_client import Client

# Инициализация клиента модели Gradio
client = Client("https://qwen-qwen2-72b-instruct.hf.space/")

# Асинхронная функция для получения данных по тендеру
async def get_tender_data(url):
    match = re.search(r'auction/(\d+)', url)
    if match:
        auction_id = match.group(1)
        print(f"Извлеченный ID тендера: {auction_id}")
    else:
        print("ID тендера не найден в URL")
        return None, None

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0 Safari/537.36',
        'Accept-Language': 'ru-RU,ru;q=0.9'
    }
    api_url = f"https://zakupki.mos.ru/newapi/api/Auction/Get?auctionId={auction_id}"
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(api_url, headers=headers)
            response.raise_for_status()
            data = response.json()

            session_info = {
                "customer": data["customer"]["name"],
                "state": data["state"]["name"],
                "startCost": data["startCost"],
                "items": [
                    {
                        "name": item["name"],
                        "currentValue": item["currentValue"],
                        "okeiName": item["okeiName"],
                        "costPerUnit": item["costPerUnit"]
                    } for item in data["items"]
                ]
            }
            files = data.get("files", [])
            file_urls = [(f"https://zakupki.mos.ru/newapi/api/FileStorage/Download?id={file['id']}", file['name']) for file in files]

            return session_info, file_urls
    except httpx.RequestError as e:
        print("Ошибка при выполнении запроса:", e)
        return None, None

# Асинхронная функция для скачивания документа
async def download_document(doc_url, save_path):
    async with httpx.AsyncClient() as client:
        response = await client.get(doc_url)
        response.raise_for_status()
        await asyncio.to_thread(save_file, save_path, response.content)
    print(f"Файл {save_path} успешно загружен.")

# Сохранение файла в отдельном потоке
def save_file(path, content):
    with open(path, 'wb') as file:
        file.write(content)

# Функция для парсинга таблиц
def parse_tables(tables):
    data_tables = {i: None for i in range(len(tables))}
    for i, table in enumerate(tables):
        data_tables[i] = [[] for _ in range(len(table.rows))]
        for j, row in enumerate(table.rows):
            for cell in row.cells:
                data_tables[i][j].append(cell.text)
    return str(data_tables)

# Асинхронная обертка для парсинга .docx
async def parse_docx(doc_path):
    return await asyncio.to_thread(_parse_docx, doc_path)

def _parse_docx(doc_path):
    doc_text = ""
    doc = Document(doc_path)
    all_tables = doc.tables
    doc_tables = parse_tables(all_tables)
    for paragraph in doc.paragraphs:
        doc_text += paragraph.text + "\n"
    return doc_text + doc_tables

# Асинхронная обертка для парсинга .pdf
async def parse_pdf(pdf_path):
    return await asyncio.to_thread(_parse_pdf, pdf_path)

def _parse_pdf(pdf_path):
    pdf_text = ""
    doc = fitz.open(pdf_path)
    for page_num in range(doc.page_count):
        page = doc[page_num]
        pdf_text += page.get_text()
    return pdf_text

# # Асинхронная функция для отправки запроса к Gradio Client
# async def generate_gradio_answer(prompt):
#     result = await asyncio.to_thread(client.predict, prompt, [["None", "None"]], "None", api_name="/model_chat")
#     return result[1][1][1]

async def predict_with_args(prompt):
    # This function wraps the client.predict call
    return await asyncio.to_thread(
        client.predict,
        prompt,                # Вводимый текст
        [["None", "None"]],   # Начальное состояние чата
        "None",               # Дополнительные параметры
        api_name="/model_chat"
    )


async def generate_gradio_answer(prompt):
    result = await predict_with_args(prompt)
    your_string = result[1][1][1]
    return your_string

# Асинхронная функция для проверки соответствия данных
async def check_compliance(session_info, doc_text):
    session_description = f"Закупка: {session_info['customer']}, начальная стоимость: {session_info['startCost']}.\n"
    session_description += "Товары:\n"
    for item in session_info['items']:
        session_description += f"- {item['name']}: {item['currentValue']} {item['okeiName']} по {item['costPerUnit']} руб.\n"

    prompt = f"""
    Проверьте соответствие спецификации закупки данным из документа.
    
    Данные из карточки тендера:
    {session_description}
    
    Данные из документа:
    {doc_text}
    
    Опишите найденные несоответствия по типу товара и количеству.
    
    Если всё в порядке, напиши "Замечаний нет"
    """
    response = await generate_gradio_answer(prompt)
    print("Ответ модели:", response)
    return response

# Функция для асинхронной обработки одного файла
async def process_file(doc, session_info):
    save_path = os.path.join("documents", doc[1])
    await download_document(doc[0], save_path)

    if save_path.endswith(".docx"):
        doc_text = await parse_docx(save_path)
    elif save_path.endswith(".pdf"):
        doc_text = await parse_pdf(save_path)
    else:
        return f"Расширение {save_path.split('.')[-1]} пока не поддерживается!"

    answer = await check_compliance(session_info, doc_text)
    return answer

# Основная асинхронная функция для обработки тендера
async def main(url):
    session_info, file_urls = await get_tender_data(url)
    if not session_info or not file_urls:
        print("Ошибка при получении данных тендера.")
        return
    
    print("Информация о тендере:", session_info)
    os.makedirs("documents", exist_ok=True)

    # Параллельная обработка всех файлов
    tasks = [process_file(doc, session_info) for doc in file_urls]
    answers = await asyncio.gather(*tasks)

    answers_str = "\n-----------------------------------------------------\n".join(answers)
    print(answers_str)
    return {
        "answers": answers_str
    }
