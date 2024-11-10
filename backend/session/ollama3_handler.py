import re
import os
import asyncio
import time
import aiohttp
import httpx  # для асинхронных запросов
from docx import Document
import fitz  # для обработки PDF
from asgiref.sync import sync_to_async
import json


# violations = "\n".join([f"- Нарушение {violation.name}" for violation in active_violations])
# print(violations)

# # Асинхронная версия функции get_active_violations
# async def get_active_violations():
#     active_violations = await sync_to_async(Violation.objects.filter)(is_active=True)
#     return "\n".join([f"- {violation.name}: {violation.example}" for violation in active_violations])

# Асинхронная функция для получения данных по тендеру
async def get_tender_data(url):
    match = re.search(r'auction/(\d+)', url)
    if match:
        auction_id = match.group(1)
        print(f"Извлеченный ID тендера: {auction_id}")
    else:
        print("ID тендера не найден в URL")
        return None, None

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0 Safari/537.36',
        'Accept-Language': 'ru-RU,ru;q=0.9'
    }
    api_url = f"https://zakupki.mos.ru/newapi/api/Auction/Get?auctionId={auction_id}"
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(api_url, headers=headers)
            response.raise_for_status()
            data = response.json()

            session_info = {
                "id": data["id"],
                "customer": data["customer"]["name"],
                "state": data["state"]["name"],
                "startCost": data["startCost"],
                "contractGuaranteeAmount": data.get("contractGuaranteeAmount") or 0, # Размер обеспечения
                "name": data["name"], # title закупки
                "federalLawName": data["federalLawName"], # "44-ФЗ",
                "conclusionReasonName": data["conclusionReasonName"], # "п. 4 ч. 1 ст. 93 Закупка объемом до 600 тысяч рублей",
                "startDate": data["startDate"], # "28.10.2024 14:22:08",
                "endDate": data["endDate"], # "29.10.2024 14:22:08",
                "termsContract": "Обязательное электронное исполнение с использованием УПД",
                "uploadLicenseDocumentsComment": data.get("uploadLicenseDocumentsComment") or "Лицензии не требуются", # "Действующая лицензия на осуществление деятельности по оказанию услуг по дезинфекции, дезинсекции и дератизации в целях обеспечения санитарно-эпидемиологического благополучия населения",
                "": data.get("deliveries", [])[0],
                "items": [
                    {
                        "name": item["name"],
                        "currentValue": item["currentValue"],
                        "okeiName": item["okeiName"],
                        "costPerUnit": item["costPerUnit"],
                        "image_url": f'https://zakupki.mos.ru/newapi/api/Core/Thumbnail/{item["imageId"]}/140/140'
                    } for item in data["items"]
                ],
                "files": data.get("files", [])
                
            }
            files = data.get("files", [])
            print("Название файла",  files[0]['name'])
            file_urls = [(f"https://zakupki.mos.ru/newapi/api/FileStorage/Download?id={file['id']}", file['name']) for file in files]

            return session_info, file_urls
    except httpx.RequestError as e:
        print("Ошибка при выполнении запроса:", e)
        return None, None

# Асинхронная функция для скачивания документа
async def download_document(doc_url, save_path):
    async with httpx.AsyncClient() as client:
        response = await client.get(doc_url)
        response.raise_for_status()
        await asyncio.to_thread(save_file, save_path, response.content)
    print(f"Файл {save_path} успешно загружен.")

# Сохранение файла в отдельном потоке
def save_file(path, content):
    with open(path, 'wb') as file:
        file.write(content)

# Функция для парсинга таблиц
def parse_tables(tables):
    data_tables = {i: None for i in range(len(tables))}
    for i, table in enumerate(tables):
        data_tables[i] = [[] for _ in range(len(table.rows))]
        for j, row in enumerate(table.rows):
            for cell in row.cells:
                data_tables[i][j].append(cell.text)
    return str(data_tables)

# Асинхронная обертка для парсинга .docx
async def parse_docx(doc_path):
    return await asyncio.to_thread(_parse_docx, doc_path)

def _parse_docx(doc_path):
    doc_text = ""
    doc = Document(doc_path)
    all_tables = doc.tables
    doc_tables = parse_tables(all_tables)
    for paragraph in doc.paragraphs:
        doc_text += paragraph.text + "\n"
    return doc_text + doc_tables

# Асинхронная обертка для парсинга .pdf
async def parse_pdf(pdf_path):
    return await asyncio.to_thread(_parse_pdf, pdf_path)

def _parse_pdf(pdf_path):
    pdf_text = ""
    doc = fitz.open(pdf_path)
    for page_num in range(doc.page_count):
        page = doc[page_num]
        pdf_text += page.get_text()
    return pdf_text

# # Асинхронная функция для отправки запроса к Gradio Client
# async def generate_gradio_answer(prompt):
#     result = await asyncio.to_thread(client.predict, prompt, [["None", "None"]], "None", api_name="/model_chat")
#     return result[1][1][1]


async def generate_detailed_response(query: str, best_match_content: str):
    url = "http://localhost:11434/api/generate"
    headers = {"Content-Type": "application/json"}
    data = {
        "model": "qwen2.5:72b",
        "prompt": f"Пользователь задал вопрос: '{query}'. Сгенерируйте подробный и полезный ответ на основании: | Содержание - {best_match_content}.",
    }

    full_response = ""  # Переменная для накопления ответа

    async with aiohttp.ClientSession() as session:
        async with session.post(url, headers=headers, json=data) as response:
            if response.status == 200:
                # Чтение NDJSON построчно
                async for line in response.content:
                    decoded_line = line.decode('utf-8').strip()  # Декодируем и очищаем строку
                    if decoded_line:
                        try:
                            # Попытка преобразовать строку в JSON
                            json_data = json.loads(decoded_line)
                            # Добавляем фрагмент ответа к полному ответу
                            full_response += json_data.get('response', '')
                            # Если ответ завершен, выходим из цикла
                            if json_data.get('done', False):
                                break
                        except json.JSONDecodeError:
                            print("Ошибка декодирования JSON:", decoded_line)
            else:
                print(f"Ошибка: {response.status}, {await response.text()}")
                return None

    print("Полный ответ:", full_response)
    return full_response


# Асинхронная функция для проверки соответствия данных
async def check_compliance(session_info, doc_text, violations):
    session_description = f"Закупка: {session_info['customer']},\n"
    session_description += f"Начальной максимальная цены контракта: {session_info['startCost']} рублей\n"

    contract_guarantee_amount = session_info['contractGuaranteeAmount']
    if contract_guarantee_amount == 0:
        session_description += f"Размер обеспечения заявки на участие в котировочной сессии: Не требуется\n"
    else:
        session_description += f"Размер обеспечения заявки на участие в котировочной сессии: {contract_guarantee_amount} рублей\n"

    session_description += f"Название котировочной сессии: {session_info['name']}\n"
    session_description += f"Требования к наличию лицензии: {session_info['uploadLicenseDocumentsComment']}\n"

    try:
        from_day = session_info['deliveries']['periodDaysFrom']
        to_day = session_info['deliveries']['periodDaysTo']
        if from_day is not None and to_day is not None:
            session_description += f"График поставки: от {from_day} до {to_day} дней\n"
        else:
            session_description += f"График поставки: c {session_info['deliveries']['periodDateFrom'].split(' ')[0]} по {session_info['deliveries']['periodDateTo'].split(' ')[0]}\n"
    except:
        session_description += f"График поставки: не найден"

    session_description += "Товары:\n"
    for item in session_info['items']:
        session_description += f"- {item['name']}: {item['currentValue']} {item['okeiName']} по {item['costPerUnit']} руб.\n"
    session_description += "Названия документов:\n"
    for doc in session_info['files']:
        session_description += f"- {doc['name']}\n"
    
    # prompt = f"""
    # Проверьте соответствие спецификации закупки данным из документов.

    # Данные из карточки закупки:
    # {session_description}

    # Данные из документов:
    # {doc_text}

    # Выявите и опишите каждый случай несоответствия по следующим критериям только если нарушение реально есть:
    # {violations} В ответе укажи только содержательную часть без повторения вопроса.
    # """
    # Формируем сессионное описание
    session_part = f"""
        Проверьте соответствие спецификации закупки данным из документов.

        Данные из карточки закупки:
        {session_description}
        """

    # Описание документов
    doc_text = doc_text.replace("\n", " ").replace("   ", " ")  # Очистка текста документа от лишних символов
    documents_part = f"Данные из документа:\n{doc_text}"

    # Критерии проверки
    criteria_part = f"""
        Выявите и опишите каждый случай несоответствия по следующим критериям, только при наличии нарушений:
        {violations}
        """

    # Описание ожидаемого формата JSON-ответа
    response_format = """
        Ответ должен быть в формате JSON и содержать следующие поля:
        - result: список нарушений, каждое из которых представляет собой объект с полями:
        - violation_type: строка, тип нарушения (например, "Несоответствие характеристик товара").
        - description: строка, описание конкретного несоответствия.
        - details: объект с полями:
            - specification: строка, данные из спецификации закупки.
            - document: строка, данные из документа.
            - difference: строка, описание разницы или причины несоответствия.
        - severity: строка, уровень серьезности ("Высокая", "Средняя", "Низкая").
        - suggestions: строка, рекомендации по исправлению.
        - summary: строка, краткое общее резюме всех нарушений.
        """

    # Объединяем части в один prompt
    prompt = session_part + documents_part + criteria_part 

    response = await generate_detailed_response(prompt, response_format)
    # print("Ответ модели:", response)
    return response

# Основная асинхронная функция для обработки тендера
async def main(url, violations):
    session_info, file_urls = await get_tender_data(url)
    if not session_info or not file_urls:
        print("Ошибка при получении данных тендера.")
        return

    violations = "\n".join([f'- Нарушение {violation["name"]}(Пример нарушения: {violation["example"]})' for violation in violations])
    print(violations)
    # print("Информация о тендере:", session_info)
    os.makedirs("documents", exist_ok=True)

    # Обработка всех файлов и объединение текстов
    answer_data = []
    
    async def process_document(doc_url, doc_name):
        combined_doc_text = ""
        save_path = os.path.join("documents", doc_name)
        await download_document(doc_url, save_path)

        if save_path.endswith(".docx"):
            combined_doc_text += await parse_docx(save_path)
        elif save_path.endswith(".pdf"):
            combined_doc_text += await parse_pdf(save_path)

        answer = await check_compliance(session_info, combined_doc_text, violations)
        
        # Обработка ответа
        try:
            answer_json = json.loads(answer)  # если ответ в JSON-строке
            return answer_json["result"]
            
        except json.JSONDecodeError:
            match = re.search(r'\{.*?\}', answer, re.DOTALL)
            if match:
                json_str = match.group(0)
                # Преобразование строки в объект JSON
                answer_json = json.loads(json_str)
                return answer_json
            else:
                return "КУПИ НОРМАЛЬНУЮ МОДЕЛЬ!"
    start_time = time.perf_counter()  # Время начала обработки

    
    # Запускаем обработку документов параллельно
    tasks = [process_document(doc_url, doc_name) for doc_url, doc_name in file_urls]
    results = await asyncio.gather(*tasks)

    # Фильтруем результаты для получения ответов
    for result in results:
        answer_data.append(result)
    end_time = time.perf_counter()  # Время окончания обработки
    processing_time = end_time - start_time
    print(f"Время обработки МОДЕЛИ!: {processing_time:.2f} секунд.")
    return {
        "answers": answer_data,
        "session_info": session_info,
        "files": [[file_url[0], file_url[1]] for file_url in file_urls]
    }